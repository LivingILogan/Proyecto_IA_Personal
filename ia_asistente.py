# Archivo: ia_asistente.py

# Se importan las librerías necesarias
import json
import sys
import os
import pyautogui
import psutil
import shutil
import re
import webbrowser
import time
import pygetwindow as gw
import pyperclip  # Importamos la librería del portapapeles
import pytesseract
from PIL import ImageGrab

# IMPORTS PARA RECONOCIMIENTO DE VOZ
import speech_recognition as sr
# Imports para investigación y web scraping
import requests
from bs4 import BeautifulSoup
# Import para crear documentos de Word
from docx import Document
from docx.shared import Inches
# IMPORTS PARA SELENIUM Y NAVEGADORES MÚLTIPLES
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException
from selenium.webdriver.firefox.service import Service as FirefoxService
from selenium.webdriver.firefox.options import Options as FirefoxOptions
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.edge.options import Options as EdgeOptions


# Nuevo: Nombre de la IA y lista de palabras de saludo y artículos
nombre_ia = "Ana"
saludos = ["hola", "qué tal", "buenos días", "buenas tardes", "buenas noches"]
articulos = ["el", "la", "los", "las", "un", "una", "unos", "unas", "al", "en"]
conjunciones = ["y", "con"]

# Mapeo de nombres comunes a nombres de archivo ejecutables
programa_map = {
    "bloc de notas": "notepad.exe",
    "bloc": "notepad.exe",
    "chrome": "chrome.exe",
    "google chrome": "chrome.exe",
    "excel": "excel.exe",
    "word": "winword.exe",
    "microsoft word": "winword.exe",
    "powerpoint": "powerpnt.exe",
    "calculadora": "calc.exe",
    "calculadora de windows": "calc.exe",
    "vscode": "code.exe",
    "visual studio code": "code.exe",
    "visual": "code.exe"
}

# Mapeo de nombres comunes a títulos de ventanas
titulo_map = {
    "bloc de notas": "bloc de notas",
    "bloc": "bloc de notas",
    "chrome": "chrome",
    "google chrome": "chrome",
    "excel": "excel",
    "word": "word",
    "microsoft word": "word",
    "powerpoint": "powerpoint",
    "calculadora": "calculadora",
    "calculadora de windows": "calculadora",
    "vscode": "visual studio code",
    "visual studio code": "visual studio code",
    "visual": "visual studio code"
}

# Diccionario para mapear rutas comunes de Windows a palabras clave
rutas_comunes = {
    "escritorio": os.path.join(os.path.expanduser('~'), 'Desktop'),
    "documentos": os.path.join(os.path.expanduser('~'), 'Documents'),
    "descargas": os.path.join(os.path.expanduser('~'), 'Downloads')
}

def guardar_lista(data):
    with open('listas.json', 'w') as file:
        json.dump(data, file, indent=4)

def leer_listas():
    try:
        with open('listas.json', 'r') as file:
            data = json.load(file)
            return data
    except FileNotFoundError:
        return []

def borrar_tareas(nombre, tareas_a_borrar):
    listas = leer_listas()
    if listas:
        for lista in listas:
            if lista.get('nombre_lista') == nombre:
                tareas_existentes = lista.get('tareas')
                tareas_actualizadas = [tarea for tarea in tareas_existentes if tarea not in tareas_a_borrar]
                lista['tareas'] = tareas_actualizadas
                guardar_lista(listas)
                return True
    return False

def guardar_interaccion(comando_usuario, accion_realizada):
    try:
        with open('datos_entrenamiento.json', 'r') as file:
            datos_existentes = json.load(file)
    except (FileNotFoundError, json.JSONDecodeError):
        datos_existentes = []
    
    nueva_interaccion = {
        'comando': comando_usuario,
        'accion': accion_realizada
    }
    datos_existentes.append(nueva_interaccion)
    
    with open('datos_entrenamiento.json', 'w') as file:
        json.dump(datos_existentes, file, indent=4)

def abrir_programa(nombre_programa):
    # Lógica de compatibilidad de plataformas
    nombre_ejecutable = programa_map.get(nombre_programa.lower(), nombre_programa)
    
    if sys.platform == 'win32':
        try:
            os.system(f'start {nombre_ejecutable}')
            time.sleep(1)
            for proc in psutil.process_iter(['name']):
                if nombre_ejecutable in proc.info['name'].lower():
                    print(f"Abriendo {nombre_programa} directamente.")
                    guardar_interaccion(f"abrir {nombre_programa}", "abrir_programa")
                    return
        except Exception as e:
            print(f"Error al intentar abrir directamente: {e}")
            
        print(f"Fallo al abrir directamente. Usando método alternativo para {nombre_programa}...")
        pyautogui.hotkey('win', 'r')
        time.sleep(1)
        pyautogui.typewrite(nombre_ejecutable)
        pyautogui.press('enter')
        print(f"Comando de 'Ejecutar' enviado para {nombre_programa}.")
        guardar_interaccion(f"abrir {nombre_programa}", "abrir_programa")
    elif sys.platform == 'darwin':
        os.system(f'open -a {nombre_programa}')
        guardar_interaccion(f"abrir {nombre_programa}", "abrir_programa")
    elif sys.platform == 'linux':
        os.system(f'xdg-open {nombre_programa}')
        guardar_interaccion(f"abrir {nombre_programa}", "abrir_programa")

def cerrar_programa_por_nombre(nombre_programa):
    for proc in psutil.process_iter():
        if nombre_programa.lower() in proc.name().lower():
            proc.kill()
            guardar_interaccion(f"cerrar {nombre_programa}", "cerrar_programa")
            return True
    return False

def crear_carpeta(nombre_carpeta):
    try:
        os.makedirs(nombre_carpeta)
        guardar_interaccion(f"crear carpeta {nombre_carpeta}", "crear_carpeta")
        return True
    except FileExistsError:
        return False

def buscar_archivo_o_carpeta(nombre):
    ruta_busqueda = "C:\\"
    print(f"Buscando '{nombre}' en la raíz del disco duro...")
    for root, dirs, files in os.walk(ruta_busqueda):
        try:
            if nombre in files:
                return os.path.join(root, nombre)
            if nombre in dirs:
                return os.path.join(root, nombre)
        except PermissionError:
            continue
    return None

def buscar_en_google(query):
    query_limpia = query.replace(' ', '+')
    url = f"https://www.google.com/search?q={query_limpia}"
    webbrowser.open(url)
    guardar_interaccion(f"buscar en google {query}", "buscar_en_google")

def guardar_playlist(canciones):
    print(f"Creando playlist con las siguientes canciones: {canciones}")
    guardar_interaccion(f"guardar playlist con {canciones}", "guardar_playlist")
    
def mover_a_escritorio(ruta_completa_a_mover):
    try:
        desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
        shutil.move(ruta_completa_a_mover, desktop_path)
        guardar_interaccion(f"mover a escritorio {ruta_completa_a_mover}", "mover_a_escritorio")
        return True
    except OSError as e:
        print(f"Error al intentar mover: {e}")
        return False

def borrar_archivo_o_carpeta(ruta_completa_a_eliminar):
    try:
        if os.path.isdir(ruta_completa_a_eliminar):
            shutil.rmtree(ruta_completa_a_eliminar)
        elif os.path.isfile(ruta_completa_a_eliminar):
            os.remove(ruta_completa_a_eliminar)
        guardar_interaccion(f"borrar {ruta_completa_a_eliminar}", "borrar_archivo_o_carpeta")
        return True
    except FileNotFoundError:
        return False
    except OSError as e:
        print(f"Error al intentar eliminar: {e}")
        return False

def hacer_clic_en_imagen(nombre_imagen):
    try:
        ubicacion = pyautogui.locateOnScreen(nombre_imagen, confidence=0.8)
        if ubicacion:
            punto_clic = pyautogui.center(ubicacion)
            pyautogui.click(punto_clic)
            print(f"Haciendo clic en la imagen '{nombre_imagen}'.")
            guardar_interaccion(f"clic en imagen {nombre_imagen}", "hacer_clic_en_imagen")
            return True
        else:
            print(f"No pude encontrar la imagen '{nombre_imagen}' en la pantalla.")
            return False
    except pyautogui.PyAutoGUIException as e:
        print(f"Error al intentar encontrar la imagen: {e}")
        return False

def activar_ventana_pyautogui(titulo_ventana):
    """
    Activa una ventana usando un bucle de espera y verificación del título.
    """
    timeout = 10 # 10 segundos para buscar
    start_time = time.time()
    
    while time.time() - start_time < timeout:
        try:
            ventanas = gw.getWindowsWithTitle(titulo_ventana)
            if ventanas:
                ventana_encontrada = ventanas[0]
                ventana_encontrada.activate()
                if gw.getActiveWindow() == ventana_encontrada:
                    print(f"Se ha activado la ventana: '{ventana_encontrada.title}'.")
                    return True
        except gw.PyGetWindowException:
            pass
            
        time.sleep(0.5)
    
    print(f"No se pudo activar la ventana '{titulo_ventana}' después de {timeout} segundos. Intenta abrirla y asegúrate de que esté visible.")
    return False

def escribir_en_programa(texto, titulo_ventana):
    """
    Copia el texto al portapapeles y lo pega en la ventana activa.
    """
    try:
        if activar_ventana_pyautogui(titulo_ventana):
            pyperclip.copy(texto)
            time.sleep(1)
            pyautogui.hotkey('ctrl', 'v')
            print("Hecho. He copiado el texto al portapapeles y lo he pegado en la ventana activa.")
            guardar_interaccion(f"escribir '{texto}' en '{titulo_ventana}'", "escribir_en_programa")
            return True
        else:
            print("No se pudo escribir el texto porque no se activó la ventana del programa.")
            return False
    except pyautogui.PyAutoGUIException as e:
        print(f"Error al intentar pegar el texto: {e}")
        return False

def abrir_y_escribir_en_programa(nombre_programa_str, texto_a_escribir):
    """
    Abre un programa, espera a que esté visible, y luego escribe en él.
    """
    nombre_programa_ejecutable = programa_map.get(nombre_programa_str)
    titulo_ventana = titulo_map.get(nombre_programa_str)

    if nombre_programa_ejecutable and titulo_ventana:
        abrir_programa(nombre_programa_str)
        print(f"Abriendo {nombre_programa_str}...")
        print(f"Esperando a que la ventana de '{titulo_ventana}' esté lista...")
        
        if not activar_ventana_pyautogui(titulo_ventana):
            print(f"No pude encontrar la ventana de '{titulo_ventana}'. Asegúrate de que el programa se abrió correctamente.")
            return

        if "word" in nombre_programa_str.lower() or "excel" in nombre_programa_str.lower():
            pyautogui.hotkey('ctrl', 'n')
            print("Hecho. He presionado Ctrl+N para crear un nuevo documento.")
            time.sleep(3)
            escribir_en_programa(texto_a_escribir, titulo_ventana)

        else:
            escribir_en_programa(texto_a_escribir, titulo_a_escribir)

        guardar_interaccion(f"abrir {nombre_programa_str} y escribir '{texto_a_escribir}'", "abrir_y_escribir")
    else:
        print("No pude encontrar el programa en mis diccionarios.")

def obtener_coordenadas_de_texto(texto_a_buscar):
    """
    Captura la pantalla, busca un texto y devuelve sus coordenadas.
    """
    try:
        # Asegúrate de que pytesseract esté configurado correctamente
        # Esta ruta puede variar en tu sistema.
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        
        # Captura de pantalla en tiempo real
        captura = ImageGrab.grab()
        
        # Usa OCR para obtener el texto y sus coordenadas
        datos_ocr = pytesseract.image_to_data(captura, output_type=pytesseract.Output.DICT)
        
        # Busca el texto en los resultados del OCR
        for i, texto in enumerate(datos_ocr['text']):
            if texto.strip().lower() == texto_a_buscar.lower():
                x = datos_ocr['left'][i]
                y = datos_ocr['top'][i]
                ancho = datos_ocr['width'][i]
                alto = datos_ocr['height'][i]
                
                # Devuelve las coordenadas centrales
                return (x + ancho // 2, y + alto // 2)

    except FileNotFoundError:
        print("Error: Tesseract no se encuentra en la ruta especificada. Por favor, revisa la ruta de instalación.")
    except Exception as e:
        print(f"Ocurrió un error en la función de visión: {e}")
    
    return None

def investigar_y_crear_documento(tema, ubicacion_cruda):
    """
    Realiza una investigación en Google sobre un tema y crea un documento de Word con formato.
    """
    print(f"Iniciando investigación sobre '{tema}'...")
    
    palabras_ubicacion = ubicacion_cruda.split()
    ubicacion = ' '.join([palabra for palabra in palabras_ubicacion if palabra not in articulos and palabra not in conjunciones])

    if ubicacion not in rutas_comunes:
        print(f"No conozco la ubicación '{ubicacion_cruda}'. Intenta con 'escritorio', 'documentos' o 'descargas'.")
        return

    informacion_relevante = {}
    
    # Lista de navegadores a intentar, en orden de preferencia
    navegadores = ['chrome', 'firefox', 'edge']

    for navegador_actual in navegadores:
        driver = None
        try:
            if navegador_actual == 'chrome':
                service = ChromeService('chromedriver.exe')
                options = ChromeOptions()
                options.add_argument("--headless")
                options.add_argument("--disable-gpu")
                options.add_argument("--log-level=3")
                driver = webdriver.Chrome(service=service, options=options)
            
            elif navegador_actual == 'firefox':
                service = FirefoxService('geckodriver.exe')
                options = FirefoxOptions()
                options.add_argument("--headless")
                driver = webdriver.Firefox(service=service, options=options)
                
            elif navegador_actual == 'edge':
                service = EdgeService('msedgedriver.exe')
                options = EdgeOptions()
                options.add_argument("--headless")
                driver = webdriver.Edge(service=service, options=options)

            print(f"✔ Conexión con {navegador_actual.capitalize()} exitosa.")
            
            # --- LÓGICA DE BÚSQUEDA Y EXTRACCIÓN DENTRO DEL BLOQUE TRY/EXCEPT ---
            query = tema.replace(' ', '+')
            google_search_url = f"https://www.google.com/search?q={query}"
            
            print(f"Navegando a Google para buscar resultados usando {navegador_actual.capitalize()}...")
            driver.get(google_search_url)
            time.sleep(2)
            
            print("Buscando enlaces de resultados de búsqueda con múltiples estrategias...")
            
            selectores_de_enlaces = [
                'h3 > a',
                'div.g a[href]',
                'a[jsname="UWckNb"]',
                'div[data-sokoban-container] a',
                'a[href^="https://www"]',
                'div.r a'
            ]
            
            search_results = []
            for selector in selectores_de_enlaces:
                try:
                    search_results = driver.find_elements(By.CSS_SELECTOR, selector)
                    if search_results:
                        print(f"✔ Estrategia '{selector}' exitosa. Se encontraron {len(search_results)} enlaces.")
                        break
                except:
                    continue
            
            if not search_results:
                print("❌ No se encontró ningún enlace con ninguna de las estrategias. Google ha cambiado su diseño por completo.")
                driver.quit()
                continue # Pasa al siguiente navegador si no se encuentran enlaces
                
            enlaces_procesados = set()
            
            for result in search_results:
                url = result.get_attribute('href')
                if url and "google.com" not in url and "youtube.com" not in url:
                    if url not in enlaces_procesados:
                        enlaces_procesados.add(url)
                        print(f"Visitando enlace: {url}")
                        
                        try:
                            driver.get(url)
                            WebDriverWait(driver, 10).until(
                                EC.presence_of_element_located((By.TAG_NAME, "body"))
                            )
                            time.sleep(3)
                            
                            soup = BeautifulSoup(driver.page_source, 'html.parser')
                            
                            main_content = soup.find('main')
                            if not main_content:
                                main_content = soup.find('article')
                            if not main_content:
                                main_content = soup.find(class_=re.compile("content|body|main|article|post|story", re.I))
                            
                            texto_pagina = ""
                            if main_content:
                                for element in main_content.find_all(['p', 'h1', 'h2', 'h3', 'li']):
                                    texto_pagina += element.get_text() + "\n\n"
                            else:
                                for element in soup.find_all(['p', 'h1', 'h2', 'h3']):
                                    texto_pagina += element.get_text() + "\n\n"
                            
                            texto_limpio = re.sub(r'[\n\s]+', ' ', texto_pagina).strip()
                            
                            if len(texto_limpio) > 300:
                                informacion_relevante[url] = texto_limpio
                                print("✔ Información relevante encontrada y guardada.")
                            else:
                                print("El contenido extraído es demasiado corto o irrelevante, se descarta.")
                            
                            if len(informacion_relevante) >= 3:
                                print("Se ha recopilado suficiente información. Saliendo de la búsqueda.")
                                break
                        except (TimeoutException, WebDriverException) as e:
                            print(f"❌ La URL {url} tardó demasiado en cargar o falló: {e}")
                            continue
            
            # Si se encontró información relevante, salimos del bucle de navegadores
            if informacion_relevante:
                break
            
        except WebDriverException as e:
            print(f"❌ Fallo al iniciar {navegador_actual.capitalize()}: {e}. Probando el siguiente...")
            if driver:
                driver.quit()
            continue # Pasa al siguiente navegador
        
        except Exception as e:
            print(f"Ocurrió un error inesperado durante la investigación: {e}")
            if driver:
                driver.quit()
            continue # Pasa al siguiente navegador
        
        finally:
            if driver:
                driver.quit()

    if not informacion_relevante:
        print("No se encontró información relevante para crear el documento.")
        return

    crear_documento_docx(tema, informacion_relevante, ubicacion)

def crear_documento_docx(titulo, contenido_dict, ubicacion):
    """
    Crea un documento de Word con formato profesional a partir de un diccionario de contenido.
    """
    document = Document()
    
    document.add_heading(f"Informe de Investigación: {titulo}", 0)
    document.add_paragraph(f"Generado por Ana el {time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    for url, texto in contenido_dict.items():
        document.add_heading(f"Fuente: {url}", level=1)
        # Limpiamos el texto para que no haya líneas en blanco excesivas
        document.add_paragraph(texto)
        
    nombre_documento = f"Informe_{titulo.replace(' ', '_')}.docx"
    ruta_destino = rutas_comunes[ubicacion]
    ruta_completa = os.path.join(ruta_destino, nombre_documento)
    
    document.save(ruta_completa)
    
    print(f"\n¡Documento de Word creado con éxito! El informe ha sido guardado como '{nombre_documento}' en la carpeta de '{ubicacion}'.")
    guardar_interaccion(f"crear documento de word '{titulo}' en '{ubicacion}'", "crear_documento_docx")

def procesar_comando(comando_procesado):
    """
    Función que procesa y ejecuta la acción del comando.
    """
    # Lógica para investigación y creación de documentos
    match_investiga = re.search(r'(?:investiga|investigar|busca|buscar) sobre (.*?) y guarda en (.*)', comando_procesado)
    if match_investiga:
        print("✔ Comando 'investiga' reconocido.")
        tema = match_investiga.group(1).strip()
        ubicacion_cruda = match_investiga.group(2).strip()
        if tema and ubicacion_cruda:
            investigar_y_crear_documento(tema, ubicacion_cruda)
        else:
            print("Por favor, dime un tema para investigar y la ubicación para guardar.")
        return

    # Lógica para guardar un documento con un nombre y ubicación específica
    match = re.search(r'guardar como (.*?) en (.*)', comando_procesado)
    if match:
        nombre_archivo = match.group(1).strip()
        ubicacion_cruda = match.group(2).strip()
        
        # Limpiar la ubicación de artículos y conjunciones
        palabras_ubicacion = ubicacion_cruda.split()
        ubicacion = ' '.join([palabra for palabra in palabras_ubicacion if palabra not in articulos and palabra not in conjunciones])

        if ubicacion in rutas_comunes:
            ruta_destino = rutas_comunes[ubicacion]
            print(f"Entendido, voy a guardar el documento '{nombre_archivo}' en la carpeta '{ubicacion}'.")

            # Si el Bloc de Notas no está abierto, lo abrimos.
            if not gw.getWindowsWithTitle("bloc de notas"):
                print("El Bloc de Notas no estaba abierto, lo estoy iniciando...")
                abrir_programa("bloc de notas")
                time.sleep(3) # Espera un momento para que se abra

            if not activar_ventana_pyautogui("bloc de notas"):
                print("No pude activar el bloc de notas para guardar el archivo.")
                return

            pyautogui.hotkey('ctrl', 'shift', 's')
            time.sleep(2)  # Pausa para que se abra el cuadro de diálogo
            
            # Escribe la ruta de la carpeta
            pyperclip.copy(ruta_destino)
            pyautogui.hotkey('ctrl', 'v')
            time.sleep(1)
            pyautogui.press('enter')
            time.sleep(1)
            
            # Escribe el nombre del archivo
            pyautogui.typewrite(nombre_archivo)
            pyautogui.press('enter')
            
            print(f"Hecho. El archivo ha sido guardado como '{nombre_archivo}' en '{ubicacion}'.")
            guardar_interaccion(comando_procesado, "guardar_como_archivo_en_ubicacion")
        else:
            print(f"No conozco la ubicación '{ubicacion_cruda}'. Intenta con 'escritorio', 'documentos' o 'descargas'.")
        return

    # Lógica para guardar un documento
    if "guardar" in comando_procesado or "salvar" in comando_procesado:
        print("Entendido, voy a guardar el documento actual.")
        
        # ACTIVA LA VENTANA ANTES DE BUSCAR
        if not activar_ventana_pyautogui("bloc de notas"):
            print("No pude activar el bloc de notas para guardar el archivo.")
            return
            
        # Encontrar y hacer clic en el menú 'Archivo'
        coordenadas_archivo = obtener_coordenadas_de_texto("Archivo")
        if coordenadas_archivo:
            pyautogui.click(coordenadas_archivo)
            time.sleep(1) # Pequeña pausa para que el menú se despliegue
            
            # Encontrar y hacer clic en 'Guardar'
            coordenadas_guardar = obtener_coordenadas_de_texto("Guardar")
            if coordenadas_guardar:
                pyautogui.click(coordenadas_guardar)
                print("Hecho. El archivo ha sido guardado.")
                guardar_interaccion(comando_procesado, "guardar_archivo")
            else:
                print("No pude encontrar la opción 'Guardar' en el menú.")
        else:
            print("No pude encontrar el menú 'Archivo'.")
        return

    # Patrones para comandos combinados
    match_abre_escribe = re.search(r'abre (.*?) y escribe (.*)', comando_procesado)
    match_escribe_en = re.search(r'escribe (.*?) en (.*)', comando_procesado)
    
    if match_abre_escribe:
        # Intención combinada: Abrir y escribir
        nombre_programa_crudo = match_abre_escribe.group(1).strip()
        nombre_programa_limpio = ' '.join([palabra for palabra in nombre_programa_crudo.split() if palabra not in articulos and palabra not in conjunciones])
        
        texto_a_escribir = match_abre_escribe.group(2).strip()
        abrir_y_escribir_en_programa(nombre_programa_limpio, texto_a_escribir)
        return

    if match_escribe_en:
        # Intención de escribir en un programa ya abierto
        texto_a_escribir = match_escribe_en.group(1).strip()
        programa_usuario_crudo = match_escribe_en.group(2).strip()
        programa_usuario_limpio = ' '.join([palabra for palabra in programa_usuario_crudo.split() if palabra not in articulos and palabra not in conjunciones])
        
        titulo_ventana = titulo_map.get(programa_usuario_limpio)
        if titulo_ventana:
            escribir_en_programa(texto_a_escribir, titulo_ventana)
        else:
            print(f"No pude encontrar el programa '{programa_usuario_limpio}' en mi diccionario de títulos.")
        return

    # Lógica para buscar y hacer clic con la visión de Ana
    if "buscar" in comando_procesado or "busca" in comando_procesado:
        # Extrae la palabra clave del comando
        partes = comando_procesado.split()
        palabras_clave = [p for p in partes if p not in ['buscar', 'busca', 'y', 'haz', 'clic', 'en', 'el', 'la', 'los', 'las']]
        
        if palabras_clave:
            texto_a_buscar = ' '.join(palabras_clave)
            coordenadas = obtener_coordenadas_de_texto(texto_a_buscar)
            
            if coordenadas:
                pyautogui.click(coordenadas)
                print(f"Hecho. Hice clic en '{texto_a_buscar}'.")
                guardar_interaccion(comando_procesado, "buscar_y_clic_en_texto")
            else:
                print(f"No pude encontrar el texto '{texto_a_buscar}' en la pantalla.")
                guardar_interaccion(comando_procesado, "texto_no_encontrado")
        else:
            print("Por favor, dime qué palabra debo buscar.")
        return

    # Intenciones simples
    palabras_comando = comando_procesado.split()
    
    if "abrir" in palabras_comando or "lanza" in palabras_comando or "ejecuta" in palabras_comando:
        indice_accion = -1
        if "abrir" in palabras_comando:
            indice_accion = palabras_comando.index("abrir")
        elif "lanza" in palabras_comando:
            indice_accion = palabras_comando.index("lanza")
        elif "ejecuta" in palabras_comando:
            indice_accion = palabras_comando.index("ejecuta")

        nombre_programa_crudo = ' '.join(palabras_comando[indice_accion+1:])
        nombre_programa_limpio = ' '.join([palabra for palabra in nombre_programa_crudo.split() if palabra not in articulos and palabra not in conjunciones])

        if nombre_programa_limpio:
            abrir_programa(nombre_programa_limpio)
            print(f"Hecho, he abierto {nombre_programa_limpio}.")
        else:
            print("No entendí qué programa quieres que abra. Por favor, sé más específico.")
        return

    # Aquí se pueden añadir las otras intenciones como "cerrar", "crear carpeta", etc.
    
    if any(keyword in comando_procesado for keyword in ["cerrar", "cierra", "detener"]):
        print("¡Claro! ¿Qué programa te gustaría cerrar?")
        nombre_programa = input("Nombre del programa: ")
        if cerrar_programa_por_nombre(nombre_programa):
            print(f"Hecho, he cerrado {nombre_programa}.")
        else:
            print(f"No pude encontrar un programa con el nombre '{nombre_programa}' para cerrar.")
        return
    
    if any(keyword in comando_procesado for keyword in ["crear", "hacer", "crea"]) and "carpeta" in comando_procesado:
        print("¡Claro! ¿Qué nombre le ponemos a la carpeta?")
        nombre_carpeta = input("Nombre de la carpeta: ")
        if crear_carpeta(nombre_carpeta):
            print(f"Hecho, he creado la carpeta '{nombre_carpeta}'.")
        else:
            print(f"La carpeta '{nombre_carpeta}' ya existe o hubo un error al crearla.")
        return

    # Funcionalidad de listas
    if any(keyword in comando_procesado for keyword in ["lista", "tarea"]) and any(keyword in comando_procesado for keyword in ["crear", "hacer", "añadir", "agregar"]):
        print("¡Entendido! Vamos a trabajar con listas.")
        if "crear" in comando_procesado or "hacer" in comando_procesado:
            nombre_lista = input("¿Qué nombre le ponemos a la lista? ")
            tareas_lista = input("¿Qué tareas quieres añadir? (separa con comas) ")
            nueva_lista = {
                'nombre_lista': nombre_lista,
                'tareas': [tarea.strip() for tarea in tareas_lista.split(',')]
            }
            listas_existentes = leer_listas()
            listas_existentes.append(nueva_lista)
            guardar_lista(listas_existentes)
            print(f"Perfecto, he creado la lista '{nombre_lista}' con tus tareas. ¡Ya no la olvidaré!")
            guardar_interaccion(f"crear lista {nombre_lista} con tareas {tareas_lista}", "crear_lista")
        elif any(keyword in comando_procesado for keyword in ["añadir", "agregar"]):
            listas = leer_listas()
            if listas:
                nombre_lista_a_modificar = input("¿A qué lista quieres añadir tareas? ")
                for lista in listas:
                    if lista.get('nombre_lista') == nombre_lista_a_modificar:
                        tareas_existentes = lista.get('tareas')
                        nuevas_tareas = input("¿Qué nuevas tareas quieres añadir? (separa con comas) ")
                        tareas_completas = tareas_existentes + [tarea.strip() for tarea in nuevas_tareas.split(',')]
                        lista['tareas'] = tareas_completas
                        guardar_lista(listas)
                        print(f"He añadido las nuevas tareas a la lista '{nombre_lista_a_modificar}'.")
                        guardar_interaccion(f"añadir tareas {nuevas_tareas} a la lista {nombre_lista_a_modificar}", "añadir_tareas_a_lista")
                        break
                else:
                    print(f"No encontré una lista con el nombre '{nombre_lista_a_modificar}'.")
            else:
                print("No tienes ninguna lista creada. Primero debes crear una.")
        elif any(keyword in comando_procesado for keyword in ["borrar", "quitar", "eliminar", "desechar"]):
            listas = leer_listas()
            if listas:
                nombre_lista_a_borrar = input("¿De qué lista quieres borrar tareas? ")
                tareas_a_borrar = input("¿Qué tareas quieres borrar? (separa con comas) ").split(',')
                if borrar_tareas(nombre_lista_a_borrar, tareas_a_borrar):
                    print(f"He borrado las tareas de la lista '{nombre_lista_a_borrar}'.")
                    guardar_interaccion(f"borrar tareas {tareas_a_borrar} de la lista {nombre_lista_a_borrar}", "borrar_tareas_de_lista")
                else:
                    print("No pude borrar las tareas. ¿Estás seguro de que la lista existe?")
            else:
                print("No tienes ninguna lista creada para borrar.")
        return

    if any(keyword in comando_procesado for keyword in ["mostrar", "ver", "revisar", "leer", "enseñar"]) and "listas" in comando_procesado:
        listas = leer_listas()
        if listas:
            print("Aquí están todas las listas que recuerdo:")
            for lista in listas:
                nombre = lista.get('nombre_lista')
                tareas = lista.get('tareas')
                print(f"\nNombre: {nombre}")
                print("Tareas:")
                for tarea in tareas:
                    print(f"- {tarea}")
            guardar_interaccion("mostrar listas", "mostrar_listas")
        else:
            print("No recuerdo ninguna lista en este momento.")
        return
        
    print("Hmm, no estoy seguro de cómo ayudarte con eso. ¿Podrías ser más específico?")
    guardar_interaccion(comando_procesado, "comando_no_reconocido")

def main():
    r = sr.Recognizer()
    try:
        # Intenta usar el micrófono primero
        with sr.Microphone() as source:
            print("Hola, soy tu IA. Estoy escuchando... 🎙️")
            r.adjust_for_ambient_noise(source)
            audio = r.listen(source)
        
        # Reconocimiento de voz en español
        entrada_usuario = r.recognize_google(audio, language="es-ES")
        print(f"Tú dijiste: {entrada_usuario}")
        
    except sr.UnknownValueError:
        print("No pude entender lo que dijiste. ¿Puedes repetirlo?")
        entrada_usuario = input("Por favor, escribe tu comando: ")
    except sr.RequestError as e:
        print(f"Error en el servicio de reconocimiento de voz; {e}")
        entrada_usuario = input("Por favor, escribe tu comando: ")
    except OSError:
        print("¡No se ha detectado un micrófono! ⚠️")
        print("Modo de entrada manual activado.")
        entrada_usuario = input("Hola, soy tu IA. ¿Qué necesitas que haga por ti? ")
    
    # Bucle principal de procesamiento de comandos
    while True:
        try:
            # Limpieza inicial del comando del usuario
            comando_procesado = entrada_usuario.lower().strip()
            
            # Limpiar el nombre de la IA y los saludos
            if comando_procesado.startswith(nombre_ia.lower()):
                comando_procesado = comando_procesado[len(nombre_ia):].strip()
            
            for saludo in saludos:
                if comando_procesado.startswith(saludo):
                    comando_procesado = comando_procesado[len(saludo):].strip()
            
            if comando_procesado in ["salir", "adiós", "bye", "terminar", "cierra la app"]:
                print("Ha sido un placer ayudarte. ¡Hasta la próxima!")
                break
            
            # Procesar y ejecutar la acción
            procesar_comando(comando_procesado)
            
            # Prepara la siguiente interacción, dependiendo de si hay micrófono
            if 'entrada_usuario' in locals():
                entrada_usuario = input("¿Algo más? ")
            else:
                try:
                    with sr.Microphone() as source:
                        print("Estoy escuchando de nuevo... 🎙️")
                        r.adjust_for_ambient_noise(source)
                        audio = r.listen(source)
                    entrada_usuario = r.recognize_google(audio, language="es-ES")
                    print(f"Tú dijiste: {entrada_usuario}")
                except (sr.UnknownValueError, sr.RequestError, OSError):
                    entrada_usuario = input("Por favor, escribe tu comando: ")
                    
        except Exception as e:
            print(f"Ocurrió un error: {e}")

if __name__ == "__main__":
    main()